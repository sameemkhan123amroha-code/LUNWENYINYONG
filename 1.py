import streamlit as st
import pandas as pd
from pypdf import PdfReader
from io import BytesIO
import openai
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.documents import Document
from sklearn.metrics.pairwise import cosine_similarity
import re
import os
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.oxml.ns import qn

# --- 1. 页面配置 ---
st.set_page_config(page_title="全能论文助手 (Pro版)", layout="wide")

st.markdown("""
<style>
    .block-container { padding-top: 20px; }
    .stButton>button { width: 100%; border-radius: 6px; height: 3em; font-weight: 600; }
    .interactive-sent { cursor: pointer; border-bottom: 2px solid #e0e0e0; padding: 0 2px; transition: all 0.2s; line-height: 1.8; }
    .interactive-sent:hover { background-color: #fff8e1; border-bottom-color: #ffc107; }
    .info-tooltip { display: none; position: fixed; background: #ffffff; border: 1px solid #d1d5da; box-shadow: 0 10px 30px rgba(0,0,0,0.15); padding: 16px; z-index: 999999; width: 400px; border-radius: 8px; font-family: sans-serif; font-size: 14px; line-height: 1.5; color: #24292e; }
    .tooltip-header { display: flex; justify-content: space-between; margin-bottom: 8px; border-bottom: 1px solid #eaecef; padding-bottom: 8px;}
    .tooltip-source { font-weight: 700; color: #0366d6; font-size: 13px; }
    .tooltip-score { font-weight: 700; font-size: 13px; }
    .tooltip-content { background: #f6f8fa; padding: 12px; border-radius: 6px; font-size: 13px; max-height: 200px; overflow-y: auto; color: #444; border: 1px solid #eaecef;}
    .spacer { height: 250px; }
</style>
""", unsafe_allow_html=True)

# --- 2. 核心逻辑函数 ---

try:
    from langchain_community.embeddings import HuggingFaceEmbeddings
except ImportError:
    HuggingFaceEmbeddings = None

def get_pdf_text_and_rename(pdf_docs, llm_renamer=None):
    """
    读取 PDF 并尝试使用 LLM 提取标题进行重命名
    """
    text_data = []
    
    # 创建进度条
    progress_bar = st.progress(0)
    status_text = st.empty()
    total_files = len(pdf_docs)

    for idx, pdf in enumerate(pdf_docs):
        try:
            status_text.text(f"正在解析第 {idx+1}/{total_files} 个文件: {pdf.name}")
            pdf_reader = PdfReader(pdf)
            text = ""
            for page in pdf_reader.pages:
                t = page.extract_text()
                if t: text += t
            
            # --- 智能重命名逻辑 ---
            final_filename = pdf.name
            if llm_renamer and len(text) > 50:
                try:
                    # 截取前 1500 个字符用于识别标题
                    sample_text = text[:1500]
                    prompt = f"任务：从以下学术论文的开头文本中提取论文标题。\n要求：直接输出标题内容，不要包含任何其他文字（如'标题是：'），不要包含文件名后缀。\n文本片段：{sample_text}"
                    res = llm_renamer.invoke(prompt)
                    new_title = res.content.strip().replace('"', '').replace('\n', '')
                    # 简单的文件名清洗，防止非法字符
                    new_title = re.sub(r'[\\/*?:"<>|]', "", new_title)
                    if len(new_title) > 2 and len(new_title) < 100: # 合理性检查
                        final_filename = f"{new_title}.pdf"
                except Exception as e:
                    print(f"重命名失败: {e}")

            text_data.append({"filename": final_filename, "text": text, "original_name": pdf.name})
            progress_bar.progress((idx + 1) / total_files)
            
        except Exception as e:
            st.error(f"⚠️ 文件 {pdf.name} 读取失败: {e}")
            
    status_text.text("解析完成！")
    progress_bar.empty()
    return text_data

def get_vectorstore(text_data, use_online_embed, api_key, api_base, provider):
    documents = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    for item in text_data:
        chunks = text_splitter.split_text(item["text"])
        for chunk in chunks:
            documents.append(Document(page_content=chunk, metadata={"source": item["filename"]}))
    
    embeddings = None
    if use_online_embed:
        if not api_key:
            st.error("❌ 使用在线 Embedding 需要提供 API Key！")
            st.stop()
        
        embed_model_name = "text-embedding-3-small"
        if "Zhipu" in provider:
            embed_model_name = "embedding-2"
        
        try:
            embeddings = OpenAIEmbeddings(
                openai_api_key=api_key, 
                openai_api_base=api_base,
                model=embed_model_name,
                chunk_size=16 
            )
        except Exception as e:
            st.error(f"❌ 在线 Embedding 初始化失败: {e}")
            st.stop()    
    else:
        if HuggingFaceEmbeddings is None:
            st.error("❌ 缺少 sentence-transformers 库")
            st.stop()
        try:
            os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'
            embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        except Exception as e:
            st.error(f"❌ 本地模型加载失败: {e}")
            st.stop()

    vectorstore = FAISS.from_documents(documents, embeddings)
    return vectorstore, embeddings

def create_word_docx(content_html, filename="综述.docx"):
    """
    将 HTML 内容转换为 Word 文档，格式：微软雅黑，小四 (12pt)
    """
    doc = DocxDocument()
    
    # 定义默认样式
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(12) # 小四 = 12pt

    # 简单清洗 HTML 标签获取纯文本 (为了 Word 格式整洁，这里只保留文本段落)
    # 如果需要保留加粗等格式，需要更复杂的 HTML 解析
    # 这里采用按段落分割的简单策略
    soup_text = re.sub(r'<[^>]+>', '\n', content_html) # 简单去标签变换行
    lines = [line.strip() for line in soup_text.split('\n') if line.strip()]

    doc.add_heading('文献综述', 0)

    for line in lines:
        p = doc.add_paragraph(line)
        # 强制设置段落字体 (有时样式继承不稳定)
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(12)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 3. 界面布局 ---

st.title("🤖 全能论文助手 (Pro版)")

# --- 配置区 ---
with st.container():
    col_config, col_upload = st.columns([1, 1.2])
    
    with col_config:
        with st.expander("🛠️ 模型参数配置", expanded=True):
            provider = st.selectbox("1. 选择大模型厂商", ["Zhipu AI (智谱GLM)", "DeepSeek (深度求索)", "OpenAI (GPT-4o)", "Kimi (月之暗面)"])
            
            p_url = "https://open.bigmodel.cn/api/paas/v4/"
            p_model = "glm-4"
            default_use_online = True 
            
            if "Zhipu" in provider:
                p_url = "https://open.bigmodel.cn/api/paas/v4/"
                p_model = "glm-4"
                default_use_online = True
            elif "DeepSeek" in provider:
                p_url = "https://api.deepseek.com/v1"
                p_model = "deepseek-chat"
                default_use_online = False
            elif "OpenAI" in provider:
                p_url = "https://api.openai.com/v1"
                p_model = "gpt-4o"
                default_use_online = True
            elif "Kimi" in provider:
                p_url = "https://api.moonshot.cn/v1"
                p_model = "moonshot-v1-8k"
                default_use_online = False
            
            api_base = st.text_input("Base URL", value=p_url)
            model_name = st.text_input("Model Name", value=p_model)
            api_key = st.text_input("API Key", type="password", placeholder="输入 Key...")
            
            st.markdown("---")
            use_online_embed = st.checkbox("使用在线 Embedding", value=default_use_online)

    with col_upload:
        uploaded_files = st.file_uploader("📂 导入 PDF 论文", accept_multiple_files=True, type=['pdf'])
        
        # 只有在有 API Key 的情况下才允许解析，因为需要用 LLM 重命名
        if uploaded_files:
            if not api_key:
                st.warning("⚠️ 请先在左侧输入 API Key，以便进行智能标题识别。")
            else:
                if 'processed_data' not in st.session_state or st.session_state.get('file_count') != len(uploaded_files):
                    if st.button("🚀 开始解析并智能重命名"):
                        # 初始化一个用于重命名的简单 LLM 实例
                        llm_renamer = ChatOpenAI(base_url=api_base, api_key=api_key, model=model_name, temperature=0.1)
                        st.session_state.processed_data = get_pdf_text_and_rename(uploaded_files, llm_renamer)
                        st.session_state.file_count = len(uploaded_files)
                        st.success(f"✅ 已加载 {len(uploaded_files)} 篇论文")
                
                # 显示解析后的文件列表
                if st.session_state.get('processed_data'):
                    with st.expander("查看已解析的论文列表"):
                        file_df = pd.DataFrame(st.session_state.processed_data)[["filename", "original_name"]]
                        st.dataframe(file_df, use_container_width=True)

st.divider()

# --- Excel 整理 ---
if st.session_state.get('processed_data'):
    col_btn, col_dl = st.columns([1, 4])
    with col_btn:
        do_excel = st.button("📊 一键整理成 EXCEL")
    
    if do_excel:
        with st.spinner(f"正在分析..."):
            try:
                llm = ChatOpenAI(base_url=api_base, api_key=api_key, model=model_name, temperature=0.1)
                summary_list = []
                prog = st.progress(0)
                total = len(st.session_state.processed_data)
                for i, item in enumerate(st.session_state.processed_data):
                    prompt = f"任务：用中文概括这篇论文的核心内容。\n【限制】：严格控制在 20 个汉字以内！直接写结论。\n论文片段：{item['text'][:2000]}"
                    res = llm.invoke(prompt)
                    # 使用新的 filename (标题)
                    summary_list.append({"论文标题 (智能识别)": item["filename"], "核心结论": res.content.strip()})
                    prog.progress((i+1)/total)
                
                df = pd.DataFrame(summary_list)
                out = BytesIO()
                with pd.ExcelWriter(out, engine='xlsxwriter') as writer:
                    df.to_excel(writer, index=False)
                    writer.sheets['Sheet1'].set_column('A:B', 40)
                
                with col_dl:
                    st.download_button("⬇️ 下载 Excel", out.getvalue(), "论文整理.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                st.dataframe(df, height=200)
            except Exception as e:
                st.error(f"API 错误: {e}")

st.divider()

# --- 综述生成 ---
st.subheader("📝 智能综述生成 (支持 Word 导出)")

c1, c2, c3 = st.columns([3, 1, 1])
query = c1.text_input("输入提示词", placeholder="例如：地聚物的抗压强度影响因素")
word_count = c2.number_input("目标字数", min_value=100, max_value=5000, value=500, step=100)
start_rag = c3.button("🚀 生成综述")

if start_rag:
    if not api_key or not st.session_state.get('processed_data') or not query:
        st.warning("⚠️ 请确保已输入 Key、上传论文并输入提示词")
    else:
        with st.spinner("🔍 检索与写作中..."):
            try:
                # 1. 检索
                vectorstore, embed_model = get_vectorstore(
                    st.session_state.processed_data, 
                    use_online_embed, 
                    api_key, 
                    api_base, 
                    provider
                )
                
                docs = vectorstore.similarity_search(query, k=8)
                if not docs:
                    st.error("未找到相关内容。")
                    st.stop()
                
                # 注意：这里 metadata['source'] 已经是修改后的标题了
                context = "\n".join([f"【来源:{d.metadata['source']}】{d.page_content}" for d in docs])
                
                # 2. 生成 (加入字数限制提示)
                prompt_text = f"""
                你是一个专业的学术助手。基于以下资料撰写关于“{query}”的综述。
                
                【写作要求】：
                1. 篇幅大约 **{word_count} 字**。
                2. 必须忠实于原文，不能编造。
                3. 在引用观点时，必须在句尾标注来源，格式为 (论文标题)。
                4. 输出格式为纯 HTML，使用 <p> 分段，不要包含 <html> 或 <body> 标签。
                
                【参考资料】：
                {context}
                """
                
                llm_chat = ChatOpenAI(base_url=api_base, api_key=api_key, model=model_name, temperature=0.3)
                resp = llm_chat.invoke(prompt_text)
                raw_html = resp.content.replace("```html", "").replace("```", "")
                
                # 3. 保存 HTML 到 session 以便导出
                st.session_state.last_generated_html = raw_html
                
                # 4. 语义核查与显示 (保持原有的高亮逻辑)
                sentences = re.split(r'(?<=[。！？])', raw_html)
                html_parts = []
                for idx, sent in enumerate(sentences):
                    clean = re.sub(r'<[^>]+>', '', sent).strip()
                    if len(clean) < 5:
                        html_parts.append(sent)
                        continue
                    
                    evidence_docs = vectorstore.similarity_search(clean, k=1)
                    if evidence_docs:
                        doc = evidence_docs[0]
                        v1 = embed_model.embed_query(clean)
                        v2 = embed_model.embed_query(doc.page_content)
                        score = cosine_similarity([v1], [v2])[0][0] * 100
                        safe_txt = doc.page_content[:300].replace('"', '&quot;').replace('\n', ' ')
                        span = f"""<span id="s_{idx}" class="interactive-sent" onclick="showTip('s_{idx}', '{doc.metadata['source']}', {round(score,1)}, '{safe_txt}')">{sent}</span>"""
                        html_parts.append(span)
                    else:
                        html_parts.append(sent)
                
                full_html = "".join(html_parts) + "<div class='spacer'></div>"
                
                # JS 代码保持不变
                js = """
                <div id="tip" class="info-tooltip">
                    <div class="tooltip-header"><span id="t-src" class="tooltip-source"></span><span id="t-score" class="tooltip-score"></span></div>
                    <div style="font-weight:bold;margin-bottom:5px">语义证据:</div>
                    <div id="t-txt" class="tooltip-content"></div>
                </div>
                <script>
                function showTip(id, src, sc, txt) {
                    var t = document.getElementById('tip');
                    var el = document.getElementById(id);
                    var r = el.getBoundingClientRect();
                    var scrollTop = window.pageYOffset || document.documentElement.scrollTop;
                    var scrollLeft = window.pageXOffset || document.documentElement.scrollLeft;
                    document.getElementById('t-src').innerText = '📄 ' + src;
                    document.getElementById('t-score').innerHTML = '匹配度: <span style="color:' + (sc>75?'#2da44e':sc>60?'#d29922':'#cf222e') + '">' + sc + '%</span>';
                    document.getElementById('t-txt').innerHTML = txt;
                    t.style.display = 'block';
                    t.style.top = (scrollTop + r.bottom + 5) + 'px';
                    t.style.left = (scrollLeft + r.left) + 'px';
                    setTimeout(() => { document.addEventListener('click', function c(e) {
                        if(e.target.id !== id && !t.contains(e.target)) { t.style.display = 'none'; document.removeEventListener('click', c); }
                    })}, 100);
                }
                </script>
                """
                
                st.success("✅ 生成完毕！")
                st.components.v1.html(f"<div style='font-family:sans-serif;padding:10px'>{full_html}</div>{js}", height=500, scrolling=True)

            except Exception as e:
                st.error(f"❌ 运行错误: {e}")

# --- Word 导出按钮 ---
if st.session_state.get('last_generated_html'):
    st.markdown("### 💾 导出结果")
    col_d1, col_d2 = st.columns([1, 4])
    with col_d1:
        # 生成 Word 文件
        docx_data = create_word_docx(st.session_state.last_generated_html)
        st.download_button(
            label="⬇️ 下载 Word 文档 (小四 微软雅黑)",
            data=docx_data,
            file_name=f"文献综述_{query}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        )

